#
#   Module responisible for document managment and printing
#

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog, simpledialog, messagebox
from docx.text.tabstops import TabStops, WD_TAB_ALIGNMENT, WD_TAB_LEADER

# extra functions
def bossChoose(toWho, doc):
    # boss choosing
    if(toWho == "pułkownik Wachulak"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("płk prof. dr hab. mgr inż. Przemysław WACHULAK")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("REKTOR-KOMENDANT")

        title_toWho2 = doc.add_paragraph()
        title_toWho2_run = title_toWho2.add_run("WOJSKOWEJ AKADEMII TECHNICZNEJ")
        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(156)

        title_toWho2_format = title_toWho2_run.font                                                                 # text format
        title_toWho2_format.name = "Times New Roman"                                                                # font name
        title_toWho2_format.size = Pt(12)                                                                           # font size (12)
        title_toWho2_format.bold = True                                                                             # bold font
        title_toWho2.paragraph_format.left_indent = Pt(156)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(156)

    if(toWho == "rektor hdk"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("REKTOR-KOMENDANT")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("WOJSKOWEJ AKADEMII")

        title_toWho2 = doc.add_paragraph()
        title_toWho2_run = title_toWho2.add_run("TECHNICZNEJ W WARSZAWIE")
        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(156)

        title_toWho2_format = title_toWho2_run.font                                                                 # text format
        title_toWho2_format.name = "Times New Roman"                                                                # font name
        title_toWho2_format.size = Pt(12)                                                                           # font size (12)
        title_toWho2_format.bold = True                                                                             # bold font
        title_toWho2.paragraph_format.left_indent = Pt(156)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(156)

# Batalies Chiefes

    if(toWho == "pułkownik Borowiecki"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("ppłk mgr inż. Konrad BOROWIECKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA BATALIONU SZKOLNEGO")
        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "pułkownik Sobecki"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("ppłk mgr inż. Grzegorz SOBECKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA BATALIONU SZKOLNEGO")
        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "pułkownik Włoch"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("ppłk mgr inż. Grzegorz WŁOCH")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA BATALIONU SZKOLNEGO")
        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

# Companies Chiefes

    if(toWho == "1"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Filip ZAWADZKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "2"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Piotr Nawalicki")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "3"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("kpt. mgr inż. Cezary CZARNECKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "4"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Albert KAROLEWSKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "5"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Bartłomiej KOTALA")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "6"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("kpt. mgr inż. Konrad WINIARSKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "7"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Andrzej PUCZKO")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "8"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Marcin ŚNIGURSKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "9"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Bogumił BORYSIEWICZ")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "10"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Aleksander STAŃKOWSKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "11"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Przemysław KAMELA")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "12"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("kpt. mgr inż. Piotr STRZELECKI")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "13"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Mateusz WRÓBEL")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "14"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("kpt. mgr inż. Paweł OSIAK")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "15"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Cezary RAMS")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "16"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Łukasz BUDNY")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)


    if(toWho == "17"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Arkadiusz KIEPAS")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "18"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Sylwia KRAUZE")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "19"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Wojciech ZARAŚ")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

    if(toWho == "20"):

        toWho_para = doc.add_paragraph()
        toWho_run = toWho_para.add_run("por. mgr inż. Jakub SZCZEŚNIAK")

        title_toWho = doc.add_paragraph()
        title_toWho_run = title_toWho.add_run("DOWÓDCA KOMPANII")

        # text settings
        title_toWho_format = title_toWho_run.font                                                                   # text format
        title_toWho_format.name = "Times New Roman"                                                                 # font name
        title_toWho_format.size = Pt(12)                                                                            # font size (12)
        title_toWho_format.bold = True                                                                              # bold font
        title_toWho.paragraph_format.left_indent = Pt(190)

        toWho_format = toWho_run.font                                                                               # text format
        toWho_format.name = "Times New Roman"                                                                       # font name
        toWho_format.size = Pt(12)                                                                                  # font size (12)
        toWho_format.bold = True                                                                                    # bold font
        toWho_para.paragraph_format.left_indent = Pt(190)

def levelBoss(toWho, doc, font_name = "Times New Roman"):
    if(toWho == "pułkownik Wachulak" or toWho == "pułkownik Borowiecki" or toWho == "pułkownik Sobecki" or toWho == "pułkownik Włoch"):
        toWho_level_var = "Pułkowniku"

        technical_text6 = doc.add_paragraph()
        technical_run6 = technical_text6.add_run("\tSzanowny Panie " + toWho_level_var + ",")                       # technical text and input toWho

        technical_format6 = technical_run6.font                                                                     # text format
        technical_format6.name = font_name                                                                          # font name
        technical_format6.size = Pt(12)                                                                             # font size (12)

    if(toWho == "3" or toWho == "6" or toWho == "12" or toWho == "14"):
        toWho_level_var = "Kapitanie"

        technical_text6 = doc.add_paragraph()
        technical_run6 = technical_text6.add_run("\tSzanowny Panie " + toWho_level_var + ",")                       # technical text and input toWho

        technical_format6 = technical_run6.font                                                                     # text format
        technical_format6.name = font_name                                                                          # font name
        technical_format6.size = Pt(12)                                                                             # font size (12)

    if(toWho == "1" or toWho == "2" or toWho == "4" or toWho == "5" or toWho == "7" or toWho == "8" or toWho == "9" or toWho == "10" or toWho == "11" or toWho == "13" or toWho == "15" or toWho == "16" or toWho == "17" or toWho == "18" or toWho == "19" or toWho == "20"):
        toWho_level_var = "Poruczniku"

        technical_text6 = doc.add_paragraph()
        technical_run6 = technical_text6.add_run("\tSzanowny Panie " + toWho_level_var + ",")                       # technical text and input toWho

        technical_format6 = technical_run6.font                                                                     # text format
        technical_format6.name = font_name                                                                          # font name
        technical_format6.size = Pt(12)                                                                             # font size (12)

def titlePosition(toWho, generation, font_name = "Times New Roman", font_size_l = Pt(12)):
    if(toWho == "pułkownik Borowiecki" or toWho == "pułkownik Sobecki" or toWho == "pułkownik Włoch" or toWho == "1" or toWho == "2" or toWho == "3" or toWho == "4" or toWho == "5" or toWho == "6" or toWho == "7" or toWho == "8" or toWho == "9" or toWho == "10" or toWho == "11" or toWho == "12" or toWho == "13" or toWho == "14" or toWho == "15" or toWho == "16" or toWho == "17" or toWho == "18" or toWho == "19" or toWho == "20"):
        technical_text3 = generation.add_paragraph()
        technical_run3 = technical_text3.add_run("WNIOSEK")                                                             # technical text
        technical_format3 = technical_run3.font                                                                         # text format
        technical_format3.name = font_name                                                                              # font name
        technical_format3.size = font_size_l                                                                            # font size (12)
        technical_format3.bold = True                                                                                   # bold font
        technical_text3.paragraph_format.left_indent = Pt(190)

    if(toWho == "pułkownik Wachulak"):
        technical_text3 = generation.add_paragraph()
        technical_run3 = technical_text3.add_run("WNIOSEK")                                                             # technical text
        technical_format3 = technical_run3.font                                                                         # text format
        technical_format3.name = font_name                                                                              # font name
        technical_format3.size = font_size_l                                                                            # font size (12)
        technical_format3.bold = True                                                                                   # bold font
        technical_text3.paragraph_format.left_indent = Pt(156)

def infoChange(topic, line, font_name = "Times New Roman", font_size_l = Pt(12)):

    if(topic == "okolicznościowy"):
        topic_var = "okolicznościowego."
        technical_run5 = line.add_run("udzielenia urlopu " + topic_var)                                             # input topic
        technical_format5 = technical_run5.font                                                                     # text format
        technical_format5.name = font_name                                                                          # font name
        technical_format5.size = font_size_l                                                                        # font size (12)
        technical_format5.italic = True                                                                             # italic font

    if(topic == "nagrodowy"):
        topic_var = "nagrodowego."
        technical_run5 = line.add_run("udzielenia urlopu " + topic_var)                                             # input topic
        technical_format5 = technical_run5.font                                                                     # text format
        technical_format5.name = font_name                                                                          # font name
        technical_format5.size = font_size_l                                                                        # font size (12)
        technical_format5.italic = True                                                                             # italic font

    if(topic == "przepustka jednorazowa"):
        topic_var = "przepustki jedorazowej."
        technical_run5 = line.add_run("udzielenia " + topic_var)                                                    # input topic
        technical_format5 = technical_run5.font                                                                     # text format
        technical_format5.name = font_name                                                                          # font name
        technical_format5.size = font_size_l                                                                        # font size (12)
        technical_format5.italic = True                                                                             # italic font

    if(topic == "HDK"):
        topic_var = "honorowego krwiodastwa."
        technical_run5 = line.add_run("zwolnienia z tytułu " + topic_var)                                             # input topic
        technical_format5 = technical_run5.font                                                                     # text format
        technical_format5.name = font_name                                                                          # font name
        technical_format5.size = font_size_l                                                                        # font size (12)
        technical_format5.italic = True                                                                             # italic font

    if(topic == "buty"):
        topic_var = "wymiany uszkodzonego obuwia."
        technical_run5 = line.add_run(topic_var)                                                                    # input topic
        technical_format5 = technical_run5.font                                                                     # text format
        technical_format5.name = font_name                                                                          # font name
        technical_format5.size = font_size_l                                                                        # font size (12)
        technical_format5.italic = True                                                                             # italic font

def check_punish(punishment, content, font_name = "Times New Roman", font_size_l = Pt(12)):

    if punishment == "brak":
        punish = "a kar dyscyplinarnych nie posiadam. "                                                              # punishment text
        punishment_run = content.add_run(punish)
        punishment_run_format = punishment_run.font                                                                  # format text
        punishment_run_format.name = font_name                                                                       # font name
        punishment_run_format.size = font_size_l                                                                     # font size (12)

    if punishment == "nagana":
        punish = "posiadam karę dyscyplinarą w wymiarze nagany. "                                                    # punishment text
        punishment_run = content.add_run(punish)
        punishment_run_format = punishment_run.font                                                                  # format text
        punishment_run_format.name = font_name                                                                       # font name
        punishment_run_format.size = font_size_l                                                                     # font size (12)

    if punishment == "5/30":
        punish = "posiadam kare dyscyplinarną w wymiarze 5/30 stawek dziennych. "                                    # punishment text
        punishment_run = content.add_run(punish)
        punishment_run_format = punishment_run.font                                                                  # format text
        punishment_run_format.name = font_name                                                                       # font name
        punishment_run_format.size = font_size_l                                                                     # font size (12)

    if punishment == "10/30":
        punish = "posiadam kare dyscyplinarną w wymiarze 10/30 stawek dziennych. "                                  # punishment text
        punishment_run = content.add_run(punish)
        punishment_run_format = punishment_run.font                                                                  # format text
        punishment_run_format.name = font_name                                                                       # font name
        punishment_run_format.size = font_size_l                                                                     # font size (12)

def check_back(back, content):
    if back != "":
        back_run = content.add_run("Posiadam zaległości z {}. ".format(back))                                       # back text
        back_format = back_run.font                                                                                 # text format
        back_format.name = "Times new Roman"                                                                        # font name
        back_format.size = Pt(12)                                                                                   # font size (12)
    else:
        back_run = content.add_run("Nie posiadam zaległości w nauce. ")                                             # back text
        back_format = back_run.font                                                                                 # text format
        back_format.name = "Times new Roman"                                                                        # font name
        back_format.size = Pt(12)                                                                                   # font size (12)

def add(add1, add2, add3, margin, generation):

    if add1 != "":
        add_para = generation.add_paragraph()
        add1_run = add_para.add_run("Załączniki: 1. " + add1 + "\n")                                                # back text
        add1_format = add1_run.font                                                                                 # text format
        add1_format.name = "Times new Roman"                                                                        # font name
        add1_format.size = Pt(10)                                                                                   # font size
        add1_margin = add_para.paragraph_format
        add1_margin.left_indent = margin                                                                            # setting left margin

        if add2 != "":
            add2_run = add_para.add_run("\t" + 10 * " " + "2. " + add2 + "\n")                                       # back text
            add2_format = add2_run.font                                                                             # text format
            add2_format.name = "Times new Roman"                                                                    # font name
            add2_format.size = Pt(10)                                                                               # font size
            add2_margin = add_para.paragraph_format
            add2_margin.left_indent = margin                                                                        # setting left margin

            if add3 != "":
                add3_run = add_para.add_run("\t" + 10 * " " + "3. " + add3 + "\n")                                   # back text
                add3_format = add3_run.font                                                                         # text format
                add3_format.name = "Times new Roman"                                                                # font name
                add3_format.size = Pt(10)                                                                           # font size
                add3_margin = add_para.paragraph_format
                add3_margin.left_indent = margin                                                                    # setting left margin
# documents
def docx_urlop(level, name, surname, where, date, group_var, direction_var, rektor, data1, data2, place, punishment, back, mot, add1, add2, add3, font_name = "Times New Roman", font_size_l = Pt(12), font_size = Pt(10), font_size_s = Pt(8), margin = Inches(0.1969)):

    file_name = simpledialog.askstring("Zapisz", "nazwa dokumentu")                                                 # window asking for name of a document
    generation = Document()                                                                                         # creating document(object)

            # filling a document
    level_n_name = generation.add_paragraph()
    level_n_name_run = level_n_name.add_run(str(level + " " + name + " " + surname.upper()))                        # printing in put level and name
    level_n_name_format = level_n_name_run.font                                                                     # format text
    level_n_name_format.name = font_name                                                                            # font name
    level_n_name_format.size = font_size                                                                            # font size (10)
    level_n_name_margin = level_n_name.paragraph_format
    level_n_name_margin.left_indent = margin                                                                        # setting left margin

    tab = level_n_name.paragraph_format.tab_stops
    tabs = tab.add_tab_stop(position = Inches(4.3), alignment = WD_TAB_ALIGNMENT.LEFT, leader = WD_TAB_LEADER.SPACES)

    place_n_date = level_n_name.add_run("\t" + where + ", dn. " + date)
    place_n_date_format = place_n_date.font                                                                         # format text
    place_n_date_format.name = font_name                                                                            # font name
    place_n_date_format.size = font_size                                                                            # font size

    technical_text = generation.add_paragraph()
    technical_run = technical_text.add_run("(stopień, imię, nazwisko podchorążego)")                                # technical text
    technical_format = technical_run.font                                                                           # format text
    technical_format.name = font_name                                                                               # font name
    technical_format.size = font_size_s                                                                             # font size (8)
    technical_text_margin = technical_text.paragraph_format
    technical_text_margin.left_indent = margin                                                                      # setting left margin

    group_n_direction = generation.add_paragraph()
    group_n_direction_run = group_n_direction.add_run(str(group_var).upper() + str(direction_var))                  # printing input group and lader
    group_n_direction_format = group_n_direction_run.font                                                           # format text
    group_n_direction_format.name = font_name                                                                       # font name
    group_n_direction_format.size = font_size                                                                       # font size (10)
    group_n_direction_margin = group_n_direction.paragraph_format
    group_n_direction_margin.left_indent = margin                                                                   # setting left margin

    technical_text2 = generation.add_paragraph()
    technical_run2 = technical_text2.add_run("(grupa studencka, pl/kp)")                                            # technical text
    technical_format2 = technical_run2.font                                                                         # text format
    technical_format2.name = font_name                                                                              # font name
    technical_format2.size = font_size_s                                                                            # font size (8)
    technical_text2_margin = technical_text2.paragraph_format
    technical_text2_margin.left_indent = margin                                                                     # setting left margin

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    bossChoose(rektor, generation)

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    titlePosition(rektor, generation)

    technical_text4 = generation.add_paragraph()
    technical_run4 = technical_text4.add_run("Dotyczy: ")                                                           # technical text
    technical_format4 = technical_run4.font                                                                         # text format
    technical_format4.name = font_name                                                                              # font name
    technical_format4.size = font_size_l                                                                            # font size (12)
    technical_format4.bold = True                                                                                   # bold font
    technical_format4.italic = True                                                                                 # italic font
    technical_text4_margin = technical_text4.paragraph_format
    technical_text4_margin.left_indent = margin                                                                     # setting left margin

    infoChange("okolicznościowy", technical_text4)

    levelBoss(rektor, generation)

    content = generation.add_paragraph()
    text_content = "zwracam się z wnioskiem o udzielenie mi urlopu {} w terminie od {} do {} Wniosek swój motywuję {}. Melduję, że w wyżej wymienionym terminie nie pełnię służb, ".format("okolicznościowego", data1, data2, mot)# content formating
    content_run = content.add_run(text_content)                                                                     # printing content
    content_format = content_run.font                                                                               # text format
    content_format.name = font_name                                                                                 # font name
    content_format.size = font_size_l                                                                               # font size (12)

    check_punish(punishment, content)
    check_back(back, content)

    text_content2 = "Na wyżej wymieniony urlop udam się do miejscowości {}.".format(place)                         # content formating
    content2_run = content.add_run(text_content2)                                                                   # printing content
    content2_format = content2_run.font                                                                             # text format
    content2_format.name = font_name                                                                                # font name
    content2_format.size = font_size_l                                                                              # font size (12)

    content_margin = content.paragraph_format
    content_margin.left_indent = margin                                                                             # setting left margin

    content.alignment =  WD_ALIGN_PARAGRAPH.JUSTIFY      								      						# position of text
    content.paragraph_format.line_spacing = 1.5                                                                     # line spacing

    please = generation.add_paragraph()
    please_run = please.add_run("\tProszę o pozytywne rozpatrzenie mojego wniosku.")                                # etiqwe run
    please_format = please_run.font                                                                                 # text format
    please_format.name = font_name                                                                                  # font name
    please_format.size = font_size_l                                                                                # font size (12)

    please_margin = please.paragraph_format
    please_margin.left_indent = margin                                                                              # setting left margin

    law = generation.add_paragraph()
    law_run = law.add_run("Podstawa: " + "\u00A7 " + "28 ust. 1 pkt. 4 Rozporządzeniem Ministra Obrony Narodowej z dn. 17 listopad 2014 r. w sprawie służby wojskowej kandydackiej na żołnierzy zawodowych.")                                  # etiqwe run
    law_format = law_run.font                                                                                       # text format
    law_format.name = font_name                                                                                     # font name
    law_format.size = font_size_l                                                                                   # font size (12)

    law_margin = law.paragraph_format
    law_margin.left_indent = margin                                                                                 # setting left margin

    null_para = generation.add_paragraph()

    technical_text7 = generation.add_paragraph()
    technical_run7 = technical_text7.add_run("z wyrazami szacunku\t\t")                                             # technical text
    technical_format7 = technical_run7.font                                                                         # text format
    technical_format7.name = font_name                                                                              # font name
    technical_format7.size = font_size_l                                                                            # font size (12)
    technical_text7.alignment =  WD_ALIGN_PARAGRAPH.RIGHT 												       		# position of text

    add(add1, add2, add3, margin, generation)

    generation.save(file_name + ".docx")														      				# saving a document

def docx_reward(level, name, surname, where, date, group_var, plut, kmp, rektor, data1, data2, place, punishment, back, nr, data, add1, add2, add3, font_name = "Times New Roman", font_size_l = Pt(12), font_size = Pt(10), font_size_s = Pt(8), margin = Inches(0.1969)):

    file_name = simpledialog.askstring("Zapisz", "nazwa dokumentu")                                                 # window asking for name of a document
    generation = Document()                                                                                         # creating document(object)

            # filling a document
    level_n_name = generation.add_paragraph()
    level_n_name_run = level_n_name.add_run(str(level + " " + name + " " + surname.upper()))                        # printing in put level and name
    level_n_name_format = level_n_name_run.font                                                                     # format text
    level_n_name_format.name = font_name                                                                            # font name
    level_n_name_format.size = font_size                                                                            # font size (10)
    level_n_name_margin = level_n_name.paragraph_format
    level_n_name_margin.left_indent = margin                                                                        # setting left margin

    tab = level_n_name.paragraph_format.tab_stops
    tabs = tab.add_tab_stop(position = Inches(4.3), alignment = WD_TAB_ALIGNMENT.LEFT, leader = WD_TAB_LEADER.SPACES)

    place_n_date = level_n_name.add_run("\t" + where + ", dn. " + date)
    place_n_date_format = place_n_date.font                                                                         # format text
    place_n_date_format.name = font_name                                                                            # font name
    place_n_date_format.size = font_size                                                                            # font size

    technical_text = generation.add_paragraph()
    technical_run = technical_text.add_run("(stopień, imię, nazwisko podchorążego)")                                # technical text
    technical_format = technical_run.font                                                                           # format text
    technical_format.name = font_name                                                                               # font name
    technical_format.size = font_size_s                                                                             # font size (8)
    technical_text_margin = technical_text.paragraph_format
    technical_text_margin.left_indent = margin                                                                      # setting left margin

    group_n_direction = generation.add_paragraph()
    group_n_direction_run = group_n_direction.add_run(str(group_var).upper() + str(plut + "pl/" + kmp + "kp"))      # printing input group and lader
    group_n_direction_format = group_n_direction_run.font                                                           # format text
    group_n_direction_format.name = font_name                                                                       # font name
    group_n_direction_format.size = font_size                                                                       # font size (10)
    group_n_direction_margin = group_n_direction.paragraph_format
    group_n_direction_margin.left_indent = margin                                                                   # setting left margin

    technical_text2 = generation.add_paragraph()
    technical_run2 = technical_text2.add_run("(grupa studencka, pl/kp)")                                            # technical text
    technical_format2 = technical_run2.font                                                                         # text format
    technical_format2.name = font_name                                                                              # font name
    technical_format2.size = font_size_s                                                                            # font size (8)
    technical_text2_margin = technical_text2.paragraph_format
    technical_text2_margin.left_indent = margin                                                                     # setting left margin

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    bossChoose(rektor, generation)

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    titlePosition(rektor, generation)

    technical_text4 = generation.add_paragraph()
    technical_run4 = technical_text4.add_run("Dotyczy: ")                                                           # technical text
    technical_format4 = technical_run4.font                                                                         # text format
    technical_format4.name = font_name                                                                              # font name
    technical_format4.size = font_size_l                                                                            # font size (12)
    technical_format4.bold = True                                                                                   # bold font
    technical_format4.italic = True                                                                                 # italic font
    technical_text4_margin = technical_text4.paragraph_format
    technical_text4_margin.left_indent = margin                                                                     # setting left margin

    infoChange("nagrodowy", technical_text4)

    levelBoss(rektor, generation)

    content = generation.add_paragraph()
    text_content = "zwracam się z wnioskiem o udzielenie mi urlopu {} w terminie od {} do {} Melduję, że w wyżej wmienionym termine nie pełnię służb, ".format("nagrodowego", data1, data2)# content formating
    content_run = content.add_run(text_content)                                                                     # printing content
    content_format = content_run.font                                                                               # text format
    content_format.name = font_name                                                                                 # font name
    content_format.size = font_size_l                                                                               # font size (12)

    check_punish(punishment, content)
    check_back(back, content)

    text_content2 = " Na wyżej wymieniony urlop udam się do miejscowości {}.".format(place)                         # content formating
    content2_run = content.add_run(text_content2)                                                                   # printing content
    content2_format = content2_run.font                                                                             # text format
    content2_format.name = font_name                                                                                # font name
    content2_format.size = font_size_l                                                                              # font size (12)

    content_margin = content.paragraph_format
    content_margin.left_indent = margin                                                                             # setting left margin

    content.alignment =  WD_ALIGN_PARAGRAPH.JUSTIFY      								      						# position of text
    content.paragraph_format.line_spacing = 1.5                                                                     # line spacing

    please = generation.add_paragraph()
    please_run = please.add_run("\tProszę o pozytywne rozpatrzenie mojego wniosku.")                                # etiqwe run
    please_format = please_run.font                                                                                 # text format
    please_format.name = font_name                                                                                  # font name
    please_format.size = font_size_l                                                                                # font size (12)

    please_margin = please.paragraph_format
    please_margin.left_indent = margin                                                                              # setting left margin

    base = generation.add_paragraph()
    text_base = "Podstawa: Rozkaz Dowódcy " + kmp + " kompanii " + "nr " + nr + " z dn. " + data + " r."             # content formating
    base_run = base.add_run(text_base)                                                                              # printing content
    base_format = base_run.font                                                                                     # text format
    base_format.name = font_name                                                                                    # font name
    base_format.size = font_size_l                                                                                  # font size (12)

    base_margin = base.paragraph_format
    base_margin.left_indent = margin                                                                                # setting left margin

    null_para = generation.add_paragraph()

    technical_text7 = generation.add_paragraph()
    technical_run7 = technical_text7.add_run("z wyrazami szacunku\t\t")                                             # technical text
    technical_format7 = technical_run7.font                                                                         # text format
    technical_format7.name = font_name                                                                              # font name
    technical_format7.size = font_size_l                                                                            # font size (12)
    technical_text7.alignment =  WD_ALIGN_PARAGRAPH.RIGHT 												       		# position of text

    add(add1, add2, add3, margin, generation)

    generation.save(file_name + ".docx")														      				# saving a document

def docx_one(level, name, surname, where, date, group_var, direction_var, kmp_var, data1, data2, place, punishment, back, mot, add1, add2, add3, font_name = "Times New Roman", font_size_l = Pt(12), font_size = Pt(10), font_size_s = Pt(8), margin = Inches(0.1969)):

    file_name = simpledialog.askstring("Zapisz", "nazwa dokumentu")                                                 # window asking for name of a document
    generation = Document()                                                                                         # creating document(object)

            # filling a document
    level_n_name = generation.add_paragraph()
    level_n_name_run = level_n_name.add_run(str(level + " " + name + " " + surname.upper()))                        # printing in put level and name
    level_n_name_format = level_n_name_run.font                                                                     # format text
    level_n_name_format.name = font_name                                                                            # font name
    level_n_name_format.size = font_size                                                                            # font size (10)
    level_n_name_margin = level_n_name.paragraph_format
    level_n_name_margin.left_indent = margin                                                                        # setting left margin

    tab = level_n_name.paragraph_format.tab_stops
    tabs = tab.add_tab_stop(position = Inches(4.3), alignment = WD_TAB_ALIGNMENT.LEFT, leader = WD_TAB_LEADER.SPACES)

    place_n_date = level_n_name.add_run("\t" + where + ", dn. " + date)
    place_n_date_format = place_n_date.font                                                                         # format text
    place_n_date_format.name = font_name                                                                            # font name
    place_n_date_format.size = font_size                                                                            # font size

    technical_text = generation.add_paragraph()
    technical_run = technical_text.add_run("(stopień, imię, nazwisko podchorążego)")                                # technical text
    technical_format = technical_run.font                                                                           # format text
    technical_format.name = font_name                                                                               # font name
    technical_format.size = font_size_s                                                                             # font size (8)
    technical_text_margin = technical_text.paragraph_format
    technical_text_margin.left_indent = margin                                                                      # setting left margin

    group_n_direction = generation.add_paragraph()
    group_n_direction_run = group_n_direction.add_run(str(group_var).upper() + str(direction_var))                  # printing input group and lader
    group_n_direction_format = group_n_direction_run.font                                                           # format text
    group_n_direction_format.name = font_name                                                                       # font name
    group_n_direction_format.size = font_size                                                                       # font size (10)
    group_n_direction_margin = group_n_direction.paragraph_format
    group_n_direction_margin.left_indent = margin                                                                   # setting left margin

    technical_text2 = generation.add_paragraph()
    technical_run2 = technical_text2.add_run("(grupa studencka, pl/kp)")                                            # technical text
    technical_format2 = technical_run2.font                                                                         # text format
    technical_format2.name = font_name                                                                              # font name
    technical_format2.size = font_size_s                                                                            # font size (8)
    technical_text2_margin = technical_text2.paragraph_format
    technical_text2_margin.left_indent = margin                                                                     # setting left margin

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    bossChoose(kmp_var, generation)

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    titlePosition(kmp_var, generation)

    technical_text4 = generation.add_paragraph()
    technical_run4 = technical_text4.add_run("Dotyczy: ")                                                           # technical text
    technical_format4 = technical_run4.font                                                                         # text format
    technical_format4.name = font_name                                                                              # font name
    technical_format4.size = font_size_l                                                                            # font size (12)
    technical_format4.bold = True                                                                                   # bold font
    technical_format4.italic = True                                                                                 # italic font
    technical_text4_margin = technical_text4.paragraph_format
    technical_text4_margin.left_indent = margin                                                                     # setting left margin

    infoChange("przepustka jednorazowa", technical_text4)

    levelBoss(kmp_var, generation)

    content = generation.add_paragraph()
    text_content = "zwracam się z wnioskiem o udzielenie mi {} w terminie od {} do {} Wniosek swój motywuję {}. Melduję, że w wyżej wymienionym terminie nie pełnię służb, ".format("przepustki jednorazowej", data1, data2, mot)# content formating
    content_run = content.add_run(text_content)                                                                     # printing content
    content_format = content_run.font                                                                               # text format
    content_format.name = font_name                                                                                 # font name
    content_format.size = font_size_l                                                                               # font size (12)

    check_punish(punishment, content)
    check_back(back, content)

    text_content2 = "Na wyżej wymieniony urlop udam się do miejscowości {}.".format(place)                          # content formating
    content2_run = content.add_run(text_content2)                                                                   # printing content
    content2_format = content2_run.font                                                                             # text format
    content2_format.name = font_name                                                                                # font name
    content2_format.size = font_size_l                                                                              # font size (12)

    content_margin = content.paragraph_format
    content_margin.left_indent = margin                                                                             # setting left margin

    content.alignment =  WD_ALIGN_PARAGRAPH.JUSTIFY      								      						# position of text
    content.paragraph_format.line_spacing = 1.5                                                                     # line spacing

    please = generation.add_paragraph()
    please_run = please.add_run("\tProszę o pozytywne rozpatrzenie mojego wniosku.")                                # etiqwe run
    please_format = please_run.font                                                                                 # text format
    please_format.name = font_name                                                                                  # font name
    please_format.size = font_size_l                                                                                # font size (12)

    please_margin = please.paragraph_format
    please_margin.left_indent = margin                                                                              # setting left margin

    null_para = generation.add_paragraph()

    technical_text7 = generation.add_paragraph()
    technical_run7 = technical_text7.add_run("z wyrazami szacunku\t\t")                                             # technical text
    technical_format7 = technical_run7.font                                                                         # text format
    technical_format7.name = font_name                                                                              # font name
    technical_format7.size = font_size_l                                                                            # font size (12)
    technical_text7.alignment =  WD_ALIGN_PARAGRAPH.RIGHT 												       		# position of text

    add(add1, add2, add3, margin, generation)

    generation.save(file_name + ".docx")														      				# saving a document

def docx_hdk(level, name, surname, where, date, group_var, direction_var, rektor, data1, data2, data3, place_var, add_var, font_name = "Times New Roman", font_size_l = Pt(12), font_size = Pt(10), font_size_s = Pt(8), margin = Inches(0.1969)):
    file_name = simpledialog.askstring("Zapisz", "nazwa dokumentu")                                                 # window asking for name of a document
    generation = Document()                                                                                         # creating document(object)

            # filling a document
    level_n_name = generation.add_paragraph()
    level_n_name_run = level_n_name.add_run(str(level + " " + name + " " + surname.upper()))                        # printing in put level and name
    level_n_name_format = level_n_name_run.font                                                                     # format text
    level_n_name_format.name = font_name                                                                            # font name
    level_n_name_format.size = font_size                                                                            # font size (10)
    level_n_name_margin = level_n_name.paragraph_format
    level_n_name_margin.left_indent = margin                                                                        # setting left margin

    tab = level_n_name.paragraph_format.tab_stops
    tabs = tab.add_tab_stop(position = Inches(4.3), alignment = WD_TAB_ALIGNMENT.LEFT, leader = WD_TAB_LEADER.SPACES)

    place_n_date = level_n_name.add_run("\t" + where + ", dn. " + date)
    place_n_date_format = place_n_date.font                                                                         # format text
    place_n_date_format.name = font_name                                                                            # font name
    place_n_date_format.size = font_size                                                                            # font size

    direction = generation.add_paragraph()
    direction_run = direction.add_run(str(direction_var))                                                           # printing input group and lader
    direction_format = direction_run.font                                                                           # format text
    direction_format.name = font_name                                                                               # font name
    direction_format.size = font_size                                                                               # font size (10)
    direction_margin = direction.paragraph_format
    direction_margin.left_indent = margin                                                                           # setting left margin

    group = generation.add_paragraph()
    group_run = group.add_run(str(group_var).upper())                                                              # printing input group and lader
    group_format = group_run.font                                                                                   # format text
    group_format.name = font_name                                                                                   # font name
    group_format.size = font_size                                                                                   # font size (10)
    group_margin = group.paragraph_format
    group_margin.left_indent = margin                                                                               # setting left margin

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    bossChoose("rektor hdk", generation)

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    technical_text = generation.add_paragraph()
    technical_run = technical_text.add_run("Dotyczy: ")                                                            # technical text
    technical_format = technical_run.font                                                                          # text format
    technical_format.name = font_name                                                                              # font name
    technical_format.size = font_size_l                                                                            # font size (12)
    technical_format.bold = True                                                                                   # bold font
    technical_format.italic = True                                                                                 # italic font
    technical_text_margin = technical_text.paragraph_format
    technical_text_margin.left_indent = margin                                                                     # setting left margin

    infoChange("HDK", technical_text)

    technical_text2 = generation.add_paragraph()
    technical_text2_run = technical_text2.add_run("\t Panie Pułkowniku, uprzejmie proszę o zgodę na zwolnienie z tytułu honorowego krwiodastwa w dniach " + str(data1) + " - " + str(data2))                                # technical text run
    technical_text2_format = technical_text2_run.font                                                               # text format
    technical_text2_format.name = font_name                                                                         # font name
    technical_text2_format.size = font_size_l                                                                       # font size (12)
    technical_text2_margin = technical_text2.paragraph_format
    technical_text2_margin.left_indent = margin                                                                     # setting left margin

    technical_text3 = generation.add_paragraph()
    technical_text3_run = technical_text3.add_run("\tMelduję, że zaległości w nauce i kar dyscyplinarnych nie mam.")# technical text run
    technical_text3_format = technical_text3_run.font                                                               # text format
    technical_text3_format.name = font_name                                                                         # font name
    technical_text3_format.size = font_size_l                                                                       # font size (12)
    technical_text3_margin = technical_text3.paragraph_format
    technical_text3_margin.left_indent = margin                                                                     # setting left margin

    technical_text4 = generation.add_paragraph()
    technical_text4_run = technical_text4.add_run("\tW w/w terminie nie pełnię służb.")                             # technical text run
    technical_text4_format = technical_text4_run.font                                                               # text format
    technical_text4_format.name = font_name                                                                         # font name
    technical_text4_format.size = font_size_l                                                                       # font size (12)
    technical_text4_margin = technical_text4.paragraph_format
    technical_text4_margin.left_indent = margin                                                                     # setting left margin

    blood_text = generation.add_paragraph()
    blood_run = blood_text.add_run("\tKrew oddałem w dniu " + data3)                                                # date run
    blood_text_format = blood_run.font                                                                              # text format
    blood_text_format.name = font_name                                                                              # font name
    blood_text_format.size = font_size_l                                                                            # font size (12)
    blood_text_margin = blood_text.paragraph_format
    blood_text_margin.left_indent = margin                                                                          # setting left margin

    place_text = generation.add_paragraph()
    place_run = place_text.add_run("\tW czasie wolnym będę przebywał w miejscowości {}.".format(place_var))         # place run
    place_text_format = place_run.font                                                                              # text format
    place_text_format.name = font_name                                                                              # font name
    place_text_format.size = font_size_l                                                                            # font size (12)
    place_text_margin =place_text.paragraph_format
    place_text_margin.left_indent = margin                                                                          # setting left margin

    please = generation.add_paragraph()
    please_run = please.add_run("\tProszę o pozytywne rozpatrzenie mojego wniosku.")                                # etiqwe run
    please_format = please_run.font                                                                                 # text format
    please_format.name = font_name                                                                                  # font name
    please_format.size = font_size_l                                                                                # font size (12)
    please_margin = please.paragraph_format
    please_margin.left_indent = margin                                                                              # setting left margin

    technical_text5 = generation.add_paragraph()
    technical_text5_run = technical_text5.add_run(" ZAŁĄCZNIKI:")                                                   # technical text run
    technical_text5_format = technical_text5_run.font                                                               # text format
    technical_text5_format.name = font_name                                                                         # font name
    technical_text5_format.size = font_size_l                                                                       # font size (12)
    technical_text5_margin = technical_text5.paragraph_format
    technical_text5_margin.left_indent = margin                                                                     # setting left margin

    add = generation.add_paragraph()
    add_run = add.add_run(add_var)
    add_format = add_run.font                                                                                       # text format
    add_format.name = font_name                                                                                     # font name
    add_format.size = font_size_l                                                                                   # font size (12)
    add_margin = add.paragraph_format
    add_margin.left_indent = margin                                                                     # setting left margin

    generation.save(file_name + ".docx")

def docx_boots(level, name, surname, where, date, group_var, direction_var, kmp_var, date_var, font_name = "Times New Roman", font_size_l = Pt(12), font_size = Pt(10), font_size_s = Pt(8), margin = Inches(0.1969)):

    file_name = simpledialog.askstring("Zapisz", "nazwa dokumentu")                                                 # window asking for name of a document
    generation = Document()                                                                                         # creating document(object)

            # filling a document
    level_n_name = generation.add_paragraph()
    level_n_name_run = level_n_name.add_run(str(level + " " + name + " " + surname.upper()))                        # printing in put level and name
    level_n_name_format = level_n_name_run.font                                                                     # format text
    level_n_name_format.name = font_name                                                                            # font name
    level_n_name_format.size = font_size                                                                            # font size (10)
    level_n_name_margin = level_n_name.paragraph_format
    level_n_name_margin.left_indent = margin                                                                        # setting left margin

    tab = level_n_name.paragraph_format.tab_stops
    tabs = tab.add_tab_stop(position = Inches(4.3), alignment = WD_TAB_ALIGNMENT.LEFT, leader = WD_TAB_LEADER.SPACES)

    place_n_date = level_n_name.add_run("\t" + where + ", dn. " + date)
    place_n_date_format = place_n_date.font                                                                         # format text
    place_n_date_format.name = font_name                                                                            # font name
    place_n_date_format.size = font_size                                                                            # font size

    technical_text = generation.add_paragraph()
    technical_run = technical_text.add_run("(stopień, imię, nazwisko podchorążego)")                                # technical text
    technical_format = technical_run.font                                                                           # format text
    technical_format.name = font_name                                                                               # font name
    technical_format.size = font_size_s                                                                             # font size (8)

    group_n_direction = generation.add_paragraph()
    group_n_direction_run = group_n_direction.add_run(str(group_var).upper() + str(direction_var))                  # printing input group and lader
    group_n_direction_format = group_n_direction_run.font                                                           # format text
    group_n_direction_format.name = font_name                                                                       # font name
    group_n_direction_format.size = font_size                                                                       # font size (10)

    technical_text2 = generation.add_paragraph()
    technical_run2 = technical_text2.add_run(7 * " " + "(grupa studencka, pl/kmp/bat)")                             # technical text
    technical_format2 = technical_run2.font                                                                         # text format
    technical_format2.name = font_name                                                                              # font name
    technical_format2.size = font_size_s                                                                            # font size (8)

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    bossChoose(kmp_var, generation)

    null_para = generation.add_paragraph()
    null_para = generation.add_paragraph()

    titlePosition(kmp_var, generation)

    technical_text4 = generation.add_paragraph()
    technical_run4 = technical_text4.add_run("Dotyczy: ")                                                           # technical text
    technical_format4 = technical_run4.font                                                                         # text format
    technical_format4.name = font_name                                                                              # font name
    technical_format4.size = font_size_l                                                                            # font size (12)
    technical_format4.bold = True                                                                                   # bold font
    technical_format4.italic = True                                                                                 # italic font

    infoChange("buty", technical_text4)

    levelBoss(kmp_var, generation)

    content = generation.add_paragraph()
    text_content = "melduje że, dnia {} podczas zajęć programowych uszkodzeniu uległo moje obuwie wojskowe. Zwracam się z prośbą o wymianę uszkodzonego obuwia na nowe.".format(date_var)# content formating
    content_run = content.add_run(text_content)                                                                     # printing content
    content_format = content_run.font                                                                               # text format
    content_format.name = font_name                                                                                 # font name
    content_format.size = font_size_l                                                                               # font size (12)

    content.alignment =  WD_ALIGN_PARAGRAPH.JUSTIFY      								      						# position of text
    content.paragraph_format.line_spacing = 1.5                                                                     # line spacing

    please = generation.add_paragraph()
    please_run = please.add_run("Proszę o pozytywne rozpatrzenie mojego wniosku.")                                  # etiqwe run
    please_format = please_run.font                                                                                 # text format
    please_format.name = font_name                                                                                  # font name
    please_format.size = font_size_l                                                                                # font size (12)

    null_para = generation.add_paragraph()

    technical_text7 = generation.add_paragraph()
    technical_run7 = technical_text7.add_run("z poważaniem\t\t")                                                    # technical text
    technical_format7 = technical_run7.font                                                                         # text format
    technical_format7.name = font_name                                                                              # font name
    technical_format7.size = font_size_l                                                                            # font size (12)
    technical_text7.alignment =  WD_ALIGN_PARAGRAPH.RIGHT 												       		# position of text

    generation.save(file_name + ".docx")														      				# saving a document
