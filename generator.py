
# -----|MODULES|----- #

from tkinter import *
from tkinter.ttk import *
from tkinter import filedialog
from tkinter import simpledialog
from datetime import datetime
from tkinter import messagebox
from os import *
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----|DATA CONVERT|----- #

def date_read():

	date = datetime.now()						# import date from local host

	if(float(date.day) < 10):					# add 0 befor day less than 10
		day = "0" + str(date.day)
	else:
		day = str(date.day)
	if(int(date.month) < 10):					# add 0 befor month less than 10
		month = "0" + str(date.month)
	else:
		month = str(date.month)

	date_right = str(day + "." + month + "." + str(date.year) + "r.")	# date formula
	return date_right

# -----|WINDOW'S SETTINGS|----- #

win = Tk()										# main window object
win.title("Generator")							# main window title
win.geometry("1200x850")						# main window geometry
frameWin = Frame(win)							# main frame of main window
frameWrite = Frame(win)							# main window's frame that add format text

# -----|TOOLBAR'S CONTROLS|----- #

def fileNew():									# option in toolbar that clear working space
	name.set("")
	place.set("")
	direct.set("")
	title.set("")
	topic.set("")
	add.set("")
	text_content.delete(0.0,"end")

def fileOpen():									# function that open file
	response = filedialog.askopenfile(initialdir = "/", title = "select", filetypes = (("documents", "*.docx"), ("all files", "*.*")))
	response = current_open_file

def fileSave():									# function that saves
	response = filedialog.asksaveasfile(initialdir = "/", title = "select", mode = "w", defaultextension = ".docx")
	response = current_open_file

def copyText():									# WORK IN PROGRESS
	text_area.clipboard_clear()
	text_area.clipboard_append(text_area.selection_get())

def pasteText():								# WORK IN PROGRESS
	text_area.insert(INSERT, text_area.clipboard_get())

def cutText():									# WORK IN PROGRESS
	text_copy()
	text_area.delete("sel.first", "sel.last")

# -----|GUI APPLICATION|----- #

screen_title = Label(frameWin, text = "Generator", font = ("arial", 27))

date = str(date_read())
date_label = Label(frameWin, text = date, font = ("arial", 10))
date_label2 = Label(frameWin, text = "Data", font = ("arial", 10))

place = StringVar()
place_entry = Entry(frameWin, textvariable = place, width = 41, font = ("arial", 10))
place_take = place_entry.get()
place_label = Label(frameWin, text = "Miejscowość: ", font = ("arial", 10))

name = StringVar()
name_entry = Entry(frameWin, textvariable = name, width = 41, font = ("arial", 10))
name_take = name_entry.get()
name_label = Label(frameWin, text = "Imię i nazwisko: ", font = ("arial", 10))

level = StringVar()
level_option = Combobox(frameWin, width = 35, textvariable = level)
level_option["values"] = ("szeregowy", "starszy szereregowy", "kapral", "starszy kapral", "plutonowy", "sierżant", "starszy sierżant")
level_option.current(0)
level_label = Label(frameWin, text = "Stopień: ", font = ("arial", 10))

toWho = StringVar()
toWho_option = Combobox(frameWin, width = 35, textvariable = toWho)
toWho_option["values"] = ("porucznik Kumosiński", "podpułkownik Borowiecki","pułkownik Wachulak")
toWho_option.current(0)
toWho_label = Label(frameWin, text = "Do kogo: ", font = ("arial", 10))

direct = StringVar()
direct_entry = Entry(frameWin, textvariable = direct, width = 41, font = ("arial", 10))
direct_take = direct_entry.get()
direct_label = Label(frameWin, text = "pluton/kompania/batalion", font = ("arial", 10))

title = StringVar()
title_entry = Entry(frameWin, textvariable = title, width = 41, font = ("arial", 10))
title_take = title_entry.get()
title_label = Label(frameWin, text = "Nagłówek: ", font = ("arial", 10))

topic = StringVar()
topic_entry = Entry(frameWin, textvariable = topic, width = 41, font = ("arial", 10))
topic_take = topic_entry.get()
topic_label = Label(frameWin, text = "Dotyczy: ", font = ("arial", 10))

add = StringVar()
add_entry = Entry(frameWin, textvariable = add, width = 41, font = ("arial", 10))
add_take = add_entry.get()
add_label = Label(frameWin, text = "załącznik: ", font = ("arial", 10))
add_button = Button(frameWin, text = "Dodaj", command = None)

text_content = Text()
text_content_label = Label(frameWrite, text = "Kontent",  font = ("arial", 12))

def acceptData():
	if name_entry.get() == "" or place_entry.get() == "" or direct_entry.get() == "" or title_entry.get() == "" or topic_entry.get() == "":	# condition to pass data
		w = messagebox.showwarning("Uwaga!", "Brak danych w wymaganych polach")										# warning for lack of data
	else:

		# -----|FILE MANAGMENT|----- #
		file_name = simpledialog.askstring("File","name")
		
		generation = Document()																						# creating new document

		level_n_name = generation.add_paragraph()																	# adding a line in document

		p = "pchor."
		if level_option.get() == "szeregowy":
			level = "   szer. " + p  + " "																			# adding a run
		if level_option.get() == "starszy szeregowy":
			level = "st. szer. " + p + " "																			# adding a run
		if level_option.get() == "kapral":
			level = "   kpr. " + p + " "																			# adding a run
		if level_option.get() == "starszy kapral":
			level = "st. kpr. " + p + " "																			# adding a run
		if level_option.get() == "plutonowy":
			level = "   plut. " + p + " "																			# adding a run
		if level_option.get() == "sierżant":
			level = "   sier. " + p + " "																			# adding a run
		if level_option.get() == "starszy sierżant":
			level = "st. sier. " + p + " "																			# adding a run

		header_run = level_n_name.add_run(level + name_entry.get())													# adding a run
		font_header_run = header_run.font																			# making a format for change
		font_header_run.name = "Times New Roman"																	# making a font
		font_header_run.size = Pt(12)																				# making a font size

		header_value = len(header_run.text)																			# making a value of text size

		header_value2 = len((place.get() + " " + date))																# making a value of text size

		header_value3 = 100 - (len(header_run.text) + header_value2)												# making a value of spaces size

		header_run3 = level_n_name.add_run(header_value3 * " ")														# adding a run for spaces
		font_header_run3 = header_run3.font																			# making a format for change
		font_header_run3.name = "Times New Roman"																	# making a font
		font_header_run3.size = Pt(12)																				# making a font size

		header_run2 = level_n_name.add_run(place.get() + " " + date)												# adding a run
		font_header_run2 = header_run2.font																			# making a format for change
		font_header_run2.name = "Times New Roman"																	# making a font
		font_header_run2.size = Pt(12)																				# making a font size

		level_n_name_label = generation.add_paragraph()																# adding a line in document
		level_n_name_label_run = level_n_name_label.add_run("	     (stopień, imię i nazwisko)")					# adding a run
		font_level_n_name_label = level_n_name_label_run.font														# making a format for adding changes
		font_level_n_name_label.name = "Times New Roman"															# making a font
		font_level_n_name_label.size = Pt(9)																		# making a size of text

		direct_label = generation.add_paragraph()																	# adding a line in document
		direct_label_run = direct_label.add_run("	    " + direct_entry.get())										# adding a run
		font_direct_label = direct_label_run.font																	# making a format for adding changes
		font_direct_label.name = "Times New Roman"																	# making a font
		font_direct_label.size = Pt(12)																				# making a size of text

		unit_label = generation.add_paragraph()																		# adding a line in document
		unit_label_run = unit_label.add_run("	(jendostka organizacyjna WAT)")										# adding a run
		font_unit_label = unit_label_run.font																		# making a format for adding changes
		font_unit_label.name = "Times New Roman"																	# making a font
		font_unit_label.size = Pt(9)																				# making a size of text

		generation.add_paragraph()
		generation.add_paragraph()

		who_para = generation.add_paragraph()																		# adding a line in document

		if(toWho.get() == "pułkownik Wachulak"):
			who_para_run = who_para.add_run("				    płk. prof. dr hab. mgr inż. Przemysław WACHULAK")	# adding a run
			font_who_para = who_para_run.font																		# making a format for adding changes
			font_who_para.name = "Times New Roman"																	# making a font
			font_who_para.size = Pt(12)																				# making a font size
			font_who_para.bold = True																				# making a font bold

			generation.add_paragraph()

			who1 = generation.add_paragraph()																		# adding a line in document
			who1_run = who1.add_run("				    REKTOR-KOMENDANT")											# adding a run
			font_who1 = who1_run.font																				# making a format for changes
			font_who1.name = "Times New Roman"																		# making a font
			font_who1.size = Pt(12)																					# making a font size
			font_who1.bold = True																					# making a font bold

			who2 = generation.add_paragraph()																		# adding a line in document
			who2_run = who2.add_run("				    WOJSKOWEJ AKADEMII TECHNICZNEJ")							# adding a run
			font_who2 = who2_run.font																				# making a format for changes
			font_who2.name = "Times New Roman"																		# making a font
			font_who2.size = Pt(12)																					# making a font size
			font_who2.bold = True																					# making a font bold

		if(toWho.get() == "podpułkownik Borowiecki"):
			who_para_run = who_para.add_run("						ppłk. mgr inż. Konrad BOROWIECKI")					# adding a run
			font_who_para = who_para_run.font																		# making a format for adding changes
			font_who_para.name = "Times New Roman"																	# making a font
			font_who_para.size = Pt(12)																				# making a font size
			font_who_para.bold = True																				# making a font bold

			generation.add_paragraph()

			who1 = generation.add_paragraph()																		# adding a line in document
			who1_run = who1.add_run("						DOWÓDCA")												# adding a run
			font_who1 = who1_run.font																				# making a format for changes
			font_who1.name = "Times New Roman"																		# making a font
			font_who1.size = Pt(12)																					# making a font size
			font_who1.bold = True																					# making a font bold

			who2 = generation.add_paragraph()																		# adding a line in document
			who2_run = who2.add_run("						1 BATALIONU SZKOLNEGO")									# adding a run
			font_who2 = who2_run.font																				# making a format for changes
			font_who2.name = "Times New Roman"																		# making a font
			font_who2.size = Pt(12)																					# making a font size
			font_who2.bold = True																					# making a font bold

		if(toWho.get() == "porucznik Kumosiński"):
			who_para_run = who_para.add_run("						por. mgr inż. Piotr KUMOSIŃSKI")				# adding a run
			font_who_para = who_para_run.font																		# making a format for adding changes
			font_who_para.name = "Times New Roman"																	# making a font
			font_who_para.size = Pt(12)																				# making a font size
			font_who_para.bold = True																				# making a font bold

			generation.add_paragraph()

			who1 = generation.add_paragraph()																		# adding a line in document
			who1_run = who1.add_run("						DOWÓDCA")												# adding a run
			font_who1 = who1_run.font																				# making a format for changes
			font_who1.name = "Times New Roman"																		# making a font
			font_who1.size = Pt(12)																					# making a font size
			font_who1.bold = True																					# making a font bold

			who2 = generation.add_paragraph()																		# adding a line in document
			who2_run = who2.add_run("						2 KOMPANII")											# adding a run
			font_who2 = who2_run.font																				# making a format for changes
			font_who2.name = "Times New Roman"																		# making a font
			font_who2.size = Pt(12)																					# making a font size
			font_who2.bold = True																					# making a font bold

		generation.add_paragraph()

		title_para =  generation.add_paragraph()																	# adding a line in document
		title_para_run = title_para.add_run(title_entry.get())														# adding a run
		font_title_para = title_para_run.font																		# making a format for adding changes
		font_title_para.name = "Times New Roman"																	# making a font
		font_title_para.size = Pt(12)																				# making a font size
		font_title_para.bold = True																					# making a font bold
		title_para.alignment =  WD_ALIGN_PARAGRAPH.CENTER 															# position of text

		generation.add_paragraph()

		topic_para = generation.add_paragraph()																		# adding a line in document
		topic_para_run = topic_para.add_run("Dotyczy: ")															# adding a run for topic label
		topic_para_run1 = topic_para.add_run(topic_entry.get())														# adding a run for topic
		font_topic_para = topic_para_run.font																		# making a format for adding changes(1st run)
		font_topic_para1 = topic_para_run1.font																		# making a format for adding changes(2nd run)
		font_topic_para.italic = True																				# making a font italic(1st run)
		font_topic_para.bold = True																					# making a font bold(1st run)
		font_topic_para.name = "Times New Roman"																	# making a font (1st run)
		font_topic_para.size = Pt(10)																				# making a font size(1st run)
		font_topic_para1.italic = True																				# making a font italic(2nd run)
		font_topic_para1.name = "Times New Roman"																	# making a font (2nd run)
		font_topic_para1.size = Pt(10)																				# making a font size(2nd run)

		generation.add_paragraph()

		content_para = generation.add_paragraph() 																	# adding a line in document
		content_para_run = content_para.add_run(text_content.get(0.0,"end"))										# adding a run
		font_content = content_para_run.font																		# making a format change
		font_content.name = "Time New Roman"																		# making a font
		font_content.size = Pt(12)																					# making a font size
		content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

		if add.get() != "":																							# if statment for adding adds
			add_para = generation.add_paragraph()																	# adding a line in document
			add_para_run = add_para.add_run("Załączniki: \n" + add_entry.get())										# adding a run
			font_add_para = add_para_run.font																		# making a format for changes
			font_add_para.name = "Times New Roman"																	# making a font
			font_add_para.size = Pt(12)																				# making a font size
		else:
			pass

		dotted = generation.add_paragraph()																			# adding a line in document
		dotted_run = dotted.add_run(".......................................................")						# adding a run
		font_dotted = dotted_run.font																				# making a format for changes
		font_dotted.name = "Times New Roman"																		# making a font
		font_dotted.size = Pt(12)																					# making a font size
		dotted.alignment =  WD_ALIGN_PARAGRAPH.RIGHT 																# position of text

		paraf = generation.add_paragraph()																			# adding a line in document
		paraf_run = paraf.add_run("(podpis)			")																# adding a run
		font_paraf = paraf_run.font																					# making a format for changes
		font_paraf.name = "Times New Roman"																			# making a font
		font_paraf.size = Pt(9)																						# making a font size
		paraf.alignment =  WD_ALIGN_PARAGRAPH.RIGHT 																# position of text


		generation.save(file_name + ".docx")																		# saving a document


		# -----|GUI RESET|----- #

		name.set("")																								# reset data from name
		place.set("")																								# reset data from place
		direct.set("")																								# reset data from direct
		title.set("")																								# reset data from title
		topic.set("")																								# reset data from topic
		add.set("")																									# reset data from add
		text_content.delete(0.0,"end")																				# reset data from content

accept_button = Button(frameWin, text = "Akceptuj", command = acceptData)

# -----|TOOLBAR'S SETTINGS|----- #

menu_obj = Menu(win)																								# menu object

file = Menu(menu_obj, tearoff = 0)																					# 'File' option (variable)
menu_obj.add_cascade(label = "File", menu = file)																	# 'File' option creation
file.add_command(label = "New", command = fileNew)																	# 'New' command
file.add_command(label = "Open", command = fileOpen)																# 'Open' command
file.add_command(label = "Save", command = fileSave)																# 'Save' command
file.add_separator()
file.add_command(label = "Exit", command = win.destroy)																# 'Exit' command

edit = Menu(menu_obj, tearoff = 0)																					# 'Edit' option (variable)
menu_obj.add_cascade(label = "Edit", menu = edit)																	# 'Edit' option creation
edit.add_command(label = "Copy", command = copyText)																# 'Copy' command
edit.add_command(label = "Cut", command = cutText)																	# 'Cut' command
edit.add_command(label = "Paste", command = pasteText)																# 'Paste' command
edit.add_command(label = "Select all", command = None)																# 'Select all' command

help = Menu(menu_obj, tearoff = 0)																					# 'Help' option (variable)
menu_obj.add_cascade(label = "Help", menu = help)																	# 'Help' option creation
help.add_command(label = "About", commend = None)																	# 'About' command
help.add_separator()
help.add_command(label = "Manual", command = None)																	# 'Manual' command

# -----|DISPLAY|----- #

screen_title.grid(column = 3, row = 0, pady = 25)

date_label2.grid(column = 1, row = 1, pady = 5)																		# actual label of date
date_label.grid(column = 2, row = 1, pady = 5)																		# label of date

place_entry.grid(column = 4, row = 1, pady = 5)																		# space for place
place_label.grid(column = 3, row = 1, pady = 5)																		# label of place

name_entry.grid(column = 2, row = 2, pady = 5)																		# space for name
name_label.grid(column = 1, row = 2, padx = 10, pady = 5)															# label of name

level_label.grid(column = 3, row = 2, pady = 5)																		# label of level
level_option.grid(column = 4, row = 2, pady = 5)																	# level choose
level_option.current()																								# default value is 1st-one

toWho_label.grid(column = 1, row = 3, pady = 5)																		# label of person
toWho_option.grid(column = 2, row = 3, padx = 10, pady = 5)															# person choose
toWho_option.current()																								# default option is 1st-one

direct_entry.grid(column = 4, row = 3, pady = 5)																	# space for direct
direct_label.grid(column = 3, row = 3, padx = 10, pady = 5)															# label of direct

title_entry.grid(column = 2, row = 4, pady = 5)																		# space for title
title_label.grid(column = 1, row = 4, pady = 5)																		# label of title

topic_entry.grid(column = 4, row = 4, pady = 5)																		# space for topic
topic_label.grid(column = 3, row = 4, pady = 5)																		# label of topic

add_entry.grid(column = 2, row = 5, pady = 5)																		# space for addjustment
add_label.grid(column = 1, row = 5, pady = 5)																		# label of addjustment
add_button.grid(column = 2, row = 6, pady = 5)																		# button that adds next addjustment

accept_button.grid(column = 5, row = 6, pady = 5)																	# button 'Akceptuj'

frameWin.pack()																										# main frame of main window
frameWrite.pack()																									# writting frame

text_content_label.pack(pady = 20)
text_content.pack(fill = BOTH, anchor = "s", padx = 285, pady = 10)													# text that is written in the main body

win.config(menu = menu_obj)																							# show toolbar
win.mainloop()																										# program loop
